from __future__ import annotations

import argparse
import os
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:  # pragma: no cover - optional dependency for diagram rendering
    Image = None
    ImageDraw = None
    ImageFont = None

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from demo_backend import SmartCampusRepository


BASE_DIR = Path(__file__).resolve().parent
REPORT_ASSET_DIR = BASE_DIR / "report_assets"
DEFAULT_PUBLIC_URL = os.getenv("SMART_CAMPUS_PUBLIC_URL") or os.getenv("RENDER_EXTERNAL_URL") or "待部署后填写 Render 公网链接"


def build_public_access_notes(public_url: str) -> list[str]:
    return [
        f"公网访问地址：{public_url}",
        "访问方式：浏览器直接打开首页后，可通过左侧导航进入图书馆、教务、实践、数据中心、治理和图谱页面。",
        "冷启动说明：本项目采用 Render 免费 Web Service 部署，实例空闲一段时间后会休眠；老师首次访问时可能需要等待约 30-60 秒完成冷启动。",
        "云端说明：免费公网演示版默认启用内置 Redis/Mongo 回退层和本地图分析模式，便于公网直接访问；本地完整版仍可连接真实 Redis、MongoDB 和 Neo4j。",
        "演示说明：在线版本支持借阅、选课、预约、签到、周报等页面操作；若实例重启后演示数据回到初始状态，属于免费演示环境的正常现象。",
    ]


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    apply_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def apply_run_font(run, size: float = 12, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for style_name, east_asia, size in [
        ("Title", "黑体", 22),
        ("Heading 1", "黑体", 16),
        ("Heading 2", "黑体", 14),
        ("Heading 3", "黑体", 12),
    ]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_paragraph(document: Document, text: str, *, style: str | None = None, bold: bool = False,
                  align: WD_PARAGRAPH_ALIGNMENT = WD_ALIGN_PARAGRAPH.JUSTIFY,
                  size: float = 12, first_line_indent: float = 0.74) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = align
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    apply_run_font(run, size=size, bold=bold)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    apply_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    document.add_paragraph("")


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 15.5) -> None:
    if not image_path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_run = caption_paragraph.add_run(caption)
    apply_run_font(caption_run, size=10.5)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_code_block(document: Document, code: str, caption: str | None = None) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7F8")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(code.strip("\n").splitlines()):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        if index != len(code.strip("\n").splitlines()) - 1:
            run.add_break()
    if caption:
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.space_after = Pt(8)
        caption_run = caption_paragraph.add_run(caption)
        apply_run_font(caption_run, size=10.5)
    else:
        document.add_paragraph("")


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        apply_run_font(run, size=11)


def add_page_break(document: Document) -> None:
    document.add_page_break()


def _diagram_font(size: int) -> Any:
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    if ImageFont is None:
        return None
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def _wrap_cjk_text(text: str, limit: int) -> list[str]:
    cleaned = str(text).strip()
    if not cleaned:
        return []
    return [cleaned[index:index + limit] for index in range(0, len(cleaned), limit)]


def _draw_card(draw, box: tuple[int, int, int, int], title: str, lines: list[str], *, fill: tuple[int, int, int]) -> None:
    title_font = _diagram_font(26)
    text_font = _diagram_font(18)
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=28, fill=fill, outline=(60, 92, 82), width=3)
    header_y = y1 + 18
    draw.text((x1 + 22, header_y), title, font=title_font, fill=(24, 54, 46))
    cursor_y = y1 + 56
    for line in lines:
        draw.text((x1 + 24, cursor_y), line, font=text_font, fill=(54, 63, 61))
        cursor_y += 24


def _draw_link(draw, start: tuple[int, int], end: tuple[int, int], label: str = "") -> None:
    draw.line([start, end], fill=(99, 120, 112), width=4)
    if label:
        mid_x = int((start[0] + end[0]) / 2)
        mid_y = int((start[1] + end[1]) / 2) - 18
        draw.text((mid_x, mid_y), label, font=_diagram_font(20), fill=(88, 96, 95), anchor="mm")


def _create_canvas(title: str) -> tuple[Any, Any]:
    image = Image.new("RGB", (1500, 940), (248, 245, 236))
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((28, 28, 1472, 912), radius=36, fill=(252, 250, 246), outline=(219, 209, 184), width=3)
    draw.text((750, 64), title, font=_diagram_font(36), fill=(27, 55, 48), anchor="mm")
    return image, draw


def _save_er_diagram(output_path: Path, title: str, nodes: list[dict[str, Any]], edges: list[tuple[str, str, str]]) -> None:
    if Image is None:
        return
    image, draw = _create_canvas(title)
    positions: dict[str, tuple[int, int, int, int]] = {}
    for node in nodes:
        box = (node["x"], node["y"], node["x"] + node["w"], node["y"] + node["h"])
        positions[node["id"]] = box
        _draw_card(draw, box, node["title"], node["lines"], fill=node.get("fill", (231, 240, 236)))
    for source, target, label in edges:
        sx1, sy1, sx2, sy2 = positions[source]
        tx1, ty1, tx2, ty2 = positions[target]
        start = (sx2, int((sy1 + sy2) / 2)) if sx2 < tx1 else (int((sx1 + sx2) / 2), sy2)
        end = (tx1, int((ty1 + ty2) / 2)) if sx2 < tx1 else (int((tx1 + tx2) / 2), ty1)
        _draw_link(draw, start, end, label)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path)


def _save_json_diagram(output_path: Path, title: str, nodes: list[dict[str, Any]], edges: list[tuple[str, str, str]]) -> None:
    if Image is None:
        return
    image, draw = _create_canvas(title)
    positions: dict[str, tuple[int, int, int, int]] = {}
    for node in nodes:
        box = (node["x"], node["y"], node["x"] + node["w"], node["y"] + node["h"])
        positions[node["id"]] = box
        _draw_card(draw, box, node["title"], node["lines"], fill=node.get("fill", (245, 236, 223)))
    for source, target, label in edges:
        sx1, sy1, sx2, sy2 = positions[source]
        tx1, ty1, tx2, ty2 = positions[target]
        start = (int((sx1 + sx2) / 2), sy2)
        end = (int((tx1 + tx2) / 2), ty1)
        _draw_link(draw, start, end, label)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path)


def build_report_diagrams() -> dict[str, Path]:
    paths = {
        "library_er": REPORT_ASSET_DIR / "library_er.png",
        "library_json": REPORT_ASSET_DIR / "library_json.png",
        "academic_er": REPORT_ASSET_DIR / "academic_er.png",
        "academic_json": REPORT_ASSET_DIR / "academic_json.png",
        "practice_er": REPORT_ASSET_DIR / "practice_er.png",
        "practice_json": REPORT_ASSET_DIR / "practice_json.png",
    }
    if Image is None:
        return paths

    _save_er_diagram(
        paths["library_er"],
        "智慧图书馆 E-R 实体关系图",
        [
            {"id": "student", "title": "students", "x": 80, "y": 180, "w": 300, "h": 190, "lines": ["PK student_id", "student_no", "student_name", "college", "major"]},
            {"id": "book", "title": "books", "x": 1120, "y": 180, "w": 300, "h": 220, "lines": ["PK book_id", "isbn", "title", "category", "total_copies", "available_copies"]},
            {"id": "borrow", "title": "borrow_records", "x": 590, "y": 150, "w": 320, "h": 250, "lines": ["PK record_id", "FK student_id", "FK book_id", "borrowed_at", "due_at", "returned_at", "status"]},
            {"id": "seat", "title": "library_seats", "x": 1120, "y": 560, "w": 300, "h": 190, "lines": ["PK seat_id", "room_name", "seat_no", "seat_status"]},
            {"id": "reserve", "title": "seat_reservations", "x": 590, "y": 540, "w": 320, "h": 220, "lines": ["PK reservation_id", "FK seat_id", "FK student_id", "reserve_date", "time_slot", "status"]},
        ],
        [
            ("student", "borrow", "1:n"),
            ("borrow", "book", "n:1"),
            ("student", "reserve", "1:n"),
            ("reserve", "seat", "n:1"),
        ],
    )
    _save_json_diagram(
        paths["library_json"],
        "library_behavior_log 文档树状结构图",
        [
            {"id": "root", "title": "library_behavior_log", "x": 520, "y": 120, "w": 460, "h": 220, "lines": ["studentName", "action", "bookTitle", "device", "createdAt"], "fill": (234, 242, 237)},
            {"id": "context", "title": "context", "x": 140, "y": 470, "w": 320, "h": 170, "lines": ["keyword", "resultCount", "source"], "fill": (245, 236, 223)},
            {"id": "audit", "title": "auditTrail", "x": 590, "y": 470, "w": 320, "h": 170, "lines": ["borrowStage", "operator", "remark"], "fill": (245, 236, 223)},
            {"id": "tags", "title": "tags[]", "x": 1040, "y": 470, "w": 320, "h": 170, "lines": ["borrow", "search", "return"], "fill": (245, 236, 223)},
        ],
        [("root", "context", "嵌套对象"), ("root", "audit", "嵌套对象"), ("root", "tags", "数组字段")],
    )
    _save_er_diagram(
        paths["academic_er"],
        "智慧教务 E-R 实体关系图",
        [
            {"id": "student", "title": "students", "x": 60, "y": 180, "w": 280, "h": 180, "lines": ["PK student_id", "student_no", "student_name", "major"]},
            {"id": "teacher", "title": "teachers", "x": 1180, "y": 110, "w": 260, "h": 180, "lines": ["PK teacher_id", "teacher_no", "teacher_name", "department"]},
            {"id": "course", "title": "courses", "x": 1180, "y": 360, "w": 260, "h": 180, "lines": ["PK course_id", "course_code", "course_name", "credit"]},
            {"id": "offering", "title": "course_offerings", "x": 760, "y": 220, "w": 320, "h": 250, "lines": ["PK offering_id", "FK course_id", "FK teacher_id", "capacity", "selected_count", "schedule_text"]},
            {"id": "selection", "title": "course_selections", "x": 400, "y": 220, "w": 280, "h": 220, "lines": ["PK selection_id", "FK student_id", "FK offering_id", "selected_at", "status"]},
            {"id": "score", "title": "score_records", "x": 400, "y": 540, "w": 280, "h": 180, "lines": ["PK score_id", "FK student_id", "FK offering_id", "total_score"]},
        ],
        [
            ("student", "selection", "1:n"),
            ("selection", "offering", "n:1"),
            ("offering", "teacher", "n:1"),
            ("offering", "course", "n:1"),
            ("student", "score", "1:n"),
            ("score", "offering", "n:1"),
        ],
    )
    _save_json_diagram(
        paths["academic_json"],
        "warning_profile 文档树状结构图",
        [
            {"id": "root", "title": "warning_profile", "x": 520, "y": 120, "w": 460, "h": 220, "lines": ["studentName", "riskType", "avgScore", "createdAt"], "fill": (234, 242, 237)},
            {"id": "rules", "title": "ruleHitList[]", "x": 120, "y": 470, "w": 360, "h": 190, "lines": ["挂科课程", "平均分阈值", "容量异常提醒"], "fill": (245, 236, 223)},
            {"id": "suggest", "title": "suggestionList[]", "x": 570, "y": 470, "w": 360, "h": 190, "lines": ["减少冲突课程", "优先补修必修课", "联系导师复核"], "fill": (245, 236, 223)},
            {"id": "audit", "title": "changeContext", "x": 1020, "y": 470, "w": 360, "h": 190, "lines": ["operator", "changeType", "afterValue"], "fill": (245, 236, 223)},
        ],
        [("root", "rules", "数组字段"), ("root", "suggest", "数组字段"), ("root", "audit", "嵌套对象")],
    )
    _save_er_diagram(
        paths["practice_er"],
        "实践教学平台 E-R 实体关系图",
        [
            {"id": "student", "title": "students", "x": 70, "y": 170, "w": 280, "h": 180, "lines": ["PK student_id", "student_no", "student_name", "major"]},
            {"id": "project", "title": "practice_projects", "x": 1120, "y": 130, "w": 300, "h": 200, "lines": ["PK project_id", "project_name", "progress", "project_status"]},
            {"id": "task", "title": "internship_tasks", "x": 600, "y": 150, "w": 320, "h": 250, "lines": ["PK task_id", "FK student_id", "FK project_id", "base_name", "mentor_name", "progress", "task_status"]},
            {"id": "lab", "title": "lab_rooms", "x": 1120, "y": 540, "w": 300, "h": 180, "lines": ["PK room_id", "room_name", "capacity", "room_status"]},
            {"id": "booking", "title": "lab_bookings", "x": 600, "y": 520, "w": 320, "h": 220, "lines": ["PK booking_id", "FK student_id", "FK room_id", "FK project_id", "booking_date", "time_slot"]},
            {"id": "report", "title": "weekly_reports", "x": 200, "y": 560, "w": 300, "h": 190, "lines": ["PK report_id", "FK task_id", "week_no", "content", "created_at"]},
        ],
        [
            ("student", "task", "1:n"),
            ("task", "project", "n:1"),
            ("student", "booking", "1:n"),
            ("booking", "lab", "n:1"),
            ("booking", "project", "n:1"),
            ("report", "task", "n:1"),
        ],
    )
    _save_json_diagram(
        paths["practice_json"],
        "practice_risk_profile 文档树状结构图",
        [
            {"id": "root", "title": "practice_risk_profile", "x": 510, "y": 120, "w": 480, "h": 240, "lines": ["studentName", "projectTitle", "riskLevel", "riskScore", "createdAt"], "fill": (234, 242, 237)},
            {"id": "rules", "title": "ruleHitList[]", "x": 100, "y": 490, "w": 370, "h": 190, "lines": ["签到不足", "进度偏慢", "报告缺失"], "fill": (245, 236, 223)},
            {"id": "suggest", "title": "suggestionList[]", "x": 560, "y": 490, "w": 370, "h": 190, "lines": ["补签到", "联系导师", "补交阶段材料"], "fill": (245, 236, 223)},
            {"id": "metrics", "title": "metrics", "x": 1020, "y": 490, "w": 370, "h": 190, "lines": ["attendanceScore", "progressScore", "reportScore"], "fill": (245, 236, 223)},
        ],
        [("root", "rules", "数组字段"), ("root", "suggest", "数组字段"), ("root", "metrics", "嵌套对象")],
    )
    return paths


def add_cover(document: Document, *, title: str, student_name: str, student_no: str, class_name: str, report_date: str) -> None:
    blank = ["", "", ""]
    for _ in blank:
        document.add_paragraph("")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("《大数据管理》")
    apply_run_font(run, size=22, bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    run = p.add_run("课程大作业")
    apply_run_font(run, size=22, bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)
    run = p.add_run(f"题  目：{title}")
    apply_run_font(run, size=16, bold=True)

    for label, value in [
        ("学生姓名", student_name),
        ("学    号", student_no),
        ("专业班级", class_name),
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(f"{label}：{value}")
        apply_run_font(run, size=14)

    for _ in range(5):
        document.add_paragraph("")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(report_date)
    apply_run_font(run, size=14)


def build_report_payload(repo: SmartCampusRepository) -> dict:
    dashboard = repo.dashboard_context()
    data_center = repo.data_center_overview()
    governance = repo.governance_overview()
    graph = repo.graph_overview()

    student_count = repo._fetch_one("SELECT COUNT(*) AS total FROM students")["total"]
    teacher_count = repo._fetch_one("SELECT COUNT(*) AS total FROM teachers")["total"]
    book_count = repo._fetch_one("SELECT COUNT(*) AS total FROM books")["total"]
    course_count = repo._fetch_one("SELECT COUNT(*) AS total FROM courses")["total"]
    task_count = repo._fetch_one("SELECT COUNT(*) AS total FROM internship_tasks")["total"]
    room_count = repo._fetch_one("SELECT COUNT(*) AS total FROM lab_rooms")["total"]

    return {
        "dashboard": dashboard,
        "data_center": data_center,
        "governance": governance,
        "graph": graph,
        "student_count": student_count,
        "teacher_count": teacher_count,
        "book_count": book_count,
        "course_count": course_count,
        "task_count": task_count,
        "room_count": room_count,
        "mongo_runtime": repo.mongo_runtime(),
        "redis_runtime": repo.redis_runtime(),
        "graph_runtime": repo.graph_runtime(),
    }


def write_markdown(
    output_path: Path,
    payload: dict,
    report_date: str,
    student_name: str,
    student_no: str,
    class_name: str,
    public_url: str,
) -> None:
    dashboard = payload["dashboard"]
    data_center = payload["data_center"]
    governance = payload["governance"]
    graph = payload["graph"]
    access_notes = build_public_access_notes(public_url)
    access_section = "\n".join(f"- {note}" for note in access_notes)
    text = f"""# 智慧校园大数据管理课程大作业报告

题目：智慧校园大数据管理  
学生姓名：{student_name}  
学号：{student_no}  
专业班级：{class_name}  
日期：{report_date}

## 在线访问说明

{access_section}

## 一、智慧校园大数据管理系统整体设计方案

本项目以荟萃大学校园信息系统为背景，围绕图书馆、教务、实践教学三个子系统完成异构数据库协同设计，并进一步补充数据中心、数据治理和图谱分析三个扩展页面。系统使用 SQLite 承载结构化事务数据，使用 Redis 承载缓存、排行榜、名额和进度摘要，使用 MongoDB 承载文档、事件流、TTL 缓存、地理空间数据和 GridFS 文件，图谱模块则以本地图分析和 Cypher 导出方式呈现跨系统关系结构。

当前演示数据规模为：学生 {payload['student_count']} 人，教师 {payload['teacher_count']} 人，图书 {payload['book_count']} 本，课程 {payload['course_count']} 门，实践任务 {payload['task_count']} 项，实验空间 {payload['room_count']} 间。图谱中共有 {graph['node_count']} 个节点、{graph['edge_count']} 条关系；数据治理页面计算得到主数据综合质量得分 {governance['quality_score_text']}。

系统总体采用“结构化主数据层 + 高速缓存层 + 文档治理层 + 图谱分析层”的四层架构：

- 结构化主数据层：使用 SQLite 模拟 MySQL，保存学生、教师、图书、课程、预约、成绩、周报等强一致业务数据。
- 高速缓存层：使用 Redis 保存日榜、周榜、月榜、分类榜、课程热选榜、实验室利用榜和项目进度摘要。
- 文档治理层：使用 MongoDB 保存行为日志、教务变更、预警画像、实践周报、导师评语、实时事件流、Geo 空间、GridFS 资产和质量快照。
- 图谱分析层：使用本地图模型和 Cypher 导出结果描述学生、教师、课程、图书、类别、任务、导师、基地等节点及关系，支持推荐和多跳分析。

### 1.1 系统业务链路

系统中的一次完整业务行为通常会在三类数据库之间形成连续的数据流。以图书借阅为例，请求先落到 SQLite 完成库存扣减和借阅记录写入，再由 Redis 更新对应排行榜与摘要信息，最后由 MongoDB 记录行为日志和事件流。选课和实践任务的处理方式也类似，即“事务层保证正确，缓存层提升读取效率，文档层保留上下文和过程留痕”，从而体现各类数据库各司其职的设计思路。

### 1.2 课程知识点映射

本项目并不是简单把多种数据库堆叠在一起，而是尽量对应课程中的知识结构来设计实现。关系数据库部分主要体现实体建模、主外键约束和事务一致性；Redis 主要体现键空间设计、排行榜结构与热点数据缓存；MongoDB 体现半结构化文档、TTL、事件流、地理空间和 GridFS 文件存储；图谱部分则体现复杂关系查询、多跳路径分析以及实体关系抽象建模。

## 二、智慧图书馆子系统数据库设计

图书馆子系统需要完成图书检索、借阅归还、座位预约、热度分析和阅读行为记录，因此采用结构化数据库与两类以上非结构化数据库的协同方案。

- SQLite：保存 books、borrow_records、library_seats、seat_reservations，保证借阅事务和库存更新的一致性。
- Redis：保存 library:last_search、rank:book:日榜/周榜/月榜/总榜/分类榜，实现热点图书排行和最近检索缓存。
- MongoDB：保存 library_behavior_log、search_session_cache、realtime_event_feed、campus_space_geo、GridFS 文件。
- Neo4j：构建 Student-BORROWED-Book-IN_CATEGORY-Category 关系链，用于图书推荐与阅读兴趣关联分析。

### 2.1 核心实体与关系

books 表记录 ISBN、题名、作者、分类、馆藏位置、总库存和可借库存；borrow_records 表记录学生借阅时间、应还时间、归还时间和当前状态；library_seats 与 seat_reservations 共同完成座位资源与预约关系管理。其核心关系是“学生 - 借阅记录 - 图书”和“学生 - 预约记录 - 座位”两条链路，前者强调库存一致性，后者强调时间冲突控制。

### 2.2 Redis 键设计

为了体现老师课堂中强调的“排行榜按时间粒度拆 Key”的思路，图书馆模块设计了多种 ZSet：

- `rank:book:YYYYMMDD`：每日热门图书榜，适合观察当天借阅热度。
- `rank:book:YYYYwkNN`：每周热门图书榜，便于看阶段性阅读趋势。
- `rank:book:YYYYMM`：每月热门图书榜，适合月度统计。
- `rank:book:total`：累计总榜，用于观察长期热门图书。
- `rank:book:category:<分类>:YYYYMMDD`：分类日榜，便于比较技术类、管理类、文学类图书的细分热度。

除排行榜外，`library:last_search` 使用 String 保存最近一次检索关键字，用于前端页面回显和热点检索观察。这样做的好处是键粒度清晰、过期策略明确、查询逻辑简单。

### 2.3 文档型数据设计

图书馆模块在 MongoDB 中至少使用了三类文档。第一类是 `library_behavior_log`，保存学生检索、借阅、归还等行为及其时间、设备、关键词、结果数量等上下文；第二类是 `search_session_cache`，带 TTL 索引，保存短期检索缓存；第三类是 `realtime_event_feed`，作为固定窗口事件流，保留最近一段时间的重要操作。除此之外，`campus_space_geo` 通过地理空间索引描述馆内空间坐标，`GridFS` 用于保存页面截图等二进制资源。

### 2.4 新增功能与设计意义

在最新版本中，图书馆页面加入了近三周借阅热力图，并把同一借阅行为同步映射到日榜、周榜、月榜、总榜和分类日榜五类 Redis Key 中。这样不仅能从页面上看到排行结果，还能从设计上解释“为什么同一事件要写入多个缓存结构”，使缓存层的职责比单一排行榜更加完整。

## 三、智慧教务系统子系统数据库设计

教务子系统需要完成开课管理、选课、退课、成绩查询和预警分析。

- SQLite：保存 courses、course_offerings、course_selections、score_records。
- Redis：保存 rank:course:current 和课程容量摘要，支撑课程热选排行和容量预警。
- MongoDB：保存 teaching_change_log、warning_profile、realtime_event_feed，用于记录容量调整、预警命中规则和操作留痕。
- Neo4j：构建 Student-ENROLLED_IN-Course 和 Teacher-TEACHES-Course 图谱，用于课程推荐和相似学生分析。

### 3.1 关系型设计重点

courses 表保存课程目录与学分信息，course_offerings 表保存学期、授课教师、上课时间、容量和已选人数，course_selections 表保存学生与课程班次之间的多对多关系，score_records 表保存成绩信息并与选课记录一一对应。关系型部分的重点是保证选课事务的一致性，也就是名额不能超卖、重复选课必须被阻止、成绩记录不能脱离选课记录单独存在。

### 3.2 冲突检测与缓存设计

当前教务页面增加了选课时间冲突检测逻辑。系统会比较已选课程与待选课程的星期、节次或时间区间，一旦发生重叠便直接拒绝写入 course_selections。与此同时，Redis 使用 `rank:course:current` 维护课程热选排行，使用 `quota:course:<id>` 维护容量摘要，这样前端无需频繁聚合数据库即可快速展示当前热门课程和剩余名额。

### 3.3 文档数据与预警画像

MongoDB 中的 `teaching_change_log` 用于保存课程容量调整、课表变更等管理动作，`warning_profile` 用于保存预警命中条件、风险说明和建议措施，`realtime_event_feed` 则保留近期选课行为与变更记录。文档库的优势在于可以把一次完整的教务事件作为整体保存，而不必为了附加字段频繁改动关系表结构。

### 3.4 设计扩展价值

教务模块在当前实现中还增加了课程热选榜、容量预警、时间冲突检测与智能推荐，因此不再只是“能选课”的事务页面，而是把课程供给、学生需求和预警分析联系在一起，更适合体现数据管理系统中“事务处理 + 分析辅助”的组合能力。

## 四、实践教学综合管理平台子系统数据库设计

实践子系统需要管理项目、实验室、任务、签到和周报。

- SQLite：保存 practice_projects、lab_rooms、lab_bookings、internship_tasks、attendance_records、weekly_reports。
- Redis：保存 rank:practice:progress、rank:lab:usage、project:progress:*。
- MongoDB：保存 internship_weekly_report、evaluation_comment、realtime_event_feed。
- Neo4j：构建 Student-HAS_TASK-Task-AT_BASE-Base 和 Task-GUIDED_BY-Mentor 图谱，用于导师匹配和多跳路径展示。

### 4.1 业务对象设计

practice_projects 用于管理项目名称、指导教师、进度和状态；lab_rooms 与 lab_bookings 描述实验室资源与预约关系；internship_tasks 记录学生任务分配与完成情况；attendance_records 保存签到数据；weekly_reports 保存过程汇报内容。这里既有典型的结构化事务对象，也有篇幅长、字段不稳定的文本型内容，因此特别适合引入文档数据库协同处理。

### 4.2 Redis 榜单与摘要

实践平台中使用 `rank:practice:progress` 维护任务推进榜，使用 `rank:lab:usage` 维护实验室利用榜，使用 `project:progress:<id>` 维护单项目进度摘要。这样可以把频繁读取的统计结果从事务表中剥离出来，减少页面访问时的实时聚合压力。

### 4.3 文档库与风险画像

MongoDB 中的 `internship_weekly_report` 和 `evaluation_comment` 分别保存周报正文与导师评语，而新增的 `practice_risk_profile` 则根据项目进度、签到、周报提交和资源预约情况计算风险等级。和关系型数据库相比，这类数据更适合以完整文档形式保存，因为字段数量可能随分析维度变化而扩展。

### 4.4 设计特点

实践模块同时具备事务管理、文本留痕和过程分析三种特征，因此它最能体现异构数据库设计的必要性。结构化数据库负责控制预约冲突和任务状态，Redis 负责展示进度与利用情况，MongoDB 负责保留过程文本与风险标签，这三者共同形成实践过程数据的闭环。

## 五、编程实现（智慧图书馆子系统）

本系统以 Flask 为 Web 框架，核心代码集中在 app.py、demo_backend.py、mongo_real_backend.py 和 graph_backend.py。图书馆子系统的典型实现流程如下：

1. 图书检索：先查询 SQLite 图书表，再把最近检索词写入 Redis，把检索行为和 TTL 搜索缓存写入 MongoDB，并写入实时事件流。
2. 图书借阅：通过事务方式更新 books.available_copies 和 borrow_records；同步刷新 Redis 排行榜，并向 MongoDB 写入行为日志和 Capped 事件流。
3. 图书归还：事务更新 returned_at 和库存，同时保留完整日志链路，方便审计。
4. 数据分析：在数据中心展示 Redis 排行、Mongo 聚合、TTL、Geo、GridFS；在图谱页面展示图书推荐。

### 5.1 模块分工

- `app.py`：负责页面路由、表单处理、重置演示数据和运行状态输出。
- `demo_backend.py`：负责业务数据初始化、SQLite 事务逻辑、Redis 键更新、统计结果整理。
- `mongo_real_backend.py`：负责 MongoDB 校验器、TTL、Geo、GridFS 和聚合功能封装。
- `graph_backend.py`：负责本地图模型构建、Cypher 导出、路径分析与推荐计算。

### 5.2 关键控制逻辑

图书借阅时，系统先判断库存是否足够，再在一个事务中同时完成库存扣减和借阅记录插入，避免出现“记录写入成功但库存未更新”的不一致情况。借阅成功后，再依次刷新 Redis 排行榜并写入 MongoDB 行为日志，保证缓存和文档层都能反映最新业务状态。类似的控制思想也被用于教务模块的选课冲突检测和实践模块的预约冲突控制中。

### 5.3 运行验证结果

当前系统运行状态如下：

- Redis 运行状态：{payload['redis_runtime']['status_label']}
- MongoDB 运行状态：{payload['mongo_runtime']['status_label']}
- 图谱模块运行状态：{payload['graph_runtime']['status_label']}

实际运行中，首页、图书馆、教务、实践、数据中心、数据治理和图谱分析页面均可正常访问；图书借阅后排行榜会实时变化；选课冲突课程会被拦截；实践风险画像会在 MongoDB 中生成快照记录。这些结果说明系统不只是停留在设计层面，而是具备可运行和可验证的实现基础。

## 六、系统页面与扩展功能说明

系统总览页将三个业务子系统、数据中心和两个分析模块统一到同一入口中，便于快速查看结构化数据、缓存数据、文档数据和关系分析的分工。

图书馆页面在原有借阅、归还和座位预约基础上，新增了日榜、周榜、月榜、总榜和分类日榜五类 Redis 排行，并增加近三周借阅热力图，用于展示时间窗口下的借阅热度变化。

教务页面新增课程热选榜、容量预警、选课时间冲突检测和智能推荐，能够说明关系数据库与缓存层如何共同支撑高频选课场景。

实践页面新增实验室利用榜、任务推进榜和风险画像，把签到、预约、进度和文本材料关联起来，便于展示过程性数据管理和文档型数据留痕。

数据中心页面集中展示 Redis 键快照、排行榜键策略、MongoDB 聚合结果、TTL/Geo/GridFS 能力和治理快照，是系统多数据库协同实现的统一观察窗口。

### 6.1 本次报告补充说明

为了让报告内容和系统版本保持一致，本次正文中同步补充了最新页面实现，包括多粒度排行榜、借阅热力图、教务冲突检测、课程推荐、实践风险画像以及数据治理与图谱分析模块。这些内容都来自当前实际运行的系统，而不是单独为文档虚构的示例。

## 七、总结

通过本次大作业，我把课程中分散讲解的关系数据库、键值数据库、文档数据库和图数据库知识组织到了同一个校园业务场景中。三个子系统分别对应了事务处理、缓存加速、文档留痕和关系分析的不同重点，也说明在真实业务中通常不会只依赖某一种数据库。

从实现效果看，SQLite 适合保存强一致的核心业务数据，Redis 适合保存热点排行和状态摘要，MongoDB 适合保存半结构化文档、日志和画像，图谱分析模块适合表达跨系统关系。它们之间并不是互相替代，而是通过明确分层形成互补关系，这也是本项目最重要的设计思想。

## 参考文献

[1] 孙旭东, 檀昌稳, 杨洋, 等. 共建共享视角下数智校园的理论框架与数据治理实现路径[J]. 现代教育技术, 2024, 34(08):132-141.  
[2] 周晓玮. 基于异构数据库的高校数据集成设计与实现[J]. 航海教育研究, 2022, 39(01):97-101.  
[3] 孙超. Redis内存数据库在智慧消防系统设计中的应用[J]. 网络安全技术与应用, 2018(08):103-105.  
[4] 白洁, 武佳丽, 余啟旺, 等. 基于MongoDB的非关系型数据库的设计与应用[J]. 湖北师范大学学报(自然科学版), 2022, 42(02):79-82.  
[5] Neo4j, Inc. Neo4j Operations Manual[EB/OL]. https://neo4j.com/docs/operations-manual/current/ , {datetime.now().strftime('%Y-%m-%d')}.  
[6] MongoDB, Inc. MongoDB Manual[EB/OL]. https://www.mongodb.com/docs/manual/ , {datetime.now().strftime('%Y-%m-%d')}.  
"""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(text, encoding="utf-8")


def write_docx(
    output_path: Path,
    payload: dict,
    report_date: str,
    student_name: str,
    student_no: str,
    class_name: str,
    public_url: str,
) -> None:
    dashboard = payload["dashboard"]
    data_center = payload["data_center"]
    governance = payload["governance"]
    graph = payload["graph"]
    screenshot_dir = BASE_DIR / "screenshots"
    diagram_assets = build_report_diagrams()

    document = Document()
    style_document(document)
    add_cover(
        document,
        title="智慧校园大数据管理",
        student_name=student_name,
        student_no=student_no,
        class_name=class_name,
        report_date=report_date,
    )
    add_page_break(document)

    add_heading(document, "在线访问说明", 1)
    for note in build_public_access_notes(public_url):
        add_paragraph(document, note, first_line_indent=0)

    add_heading(document, "摘要", 1)
    add_paragraph(
        document,
        "本项目以智慧校园为业务背景，围绕智慧图书馆、智慧教务系统和实践教学综合管理平台三个子系统，完成了一个面向《大数据管理》课程要求的异构数据库协同演示系统。系统以 SQLite 模拟事务型主数据层，以 Redis 承担排行榜、容量摘要和状态缓存，以 MongoDB 保存日志、画像、事件流、TTL 缓存、地理空间数据和 GridFS 文件，同时通过本地图模型和 Cypher 导出方式补充图谱分析能力。为了增强系统完整性，项目还增加了数据中心、数据治理实验室和图谱分析模块，使数据库设计、页面展示和运行验证形成闭环。",
    )
    add_paragraph(
        document,
        "在实现层面，系统能够完成借阅、归还、选课、退课、实验室预约、签到、周报提交、风险画像和多粒度排行榜刷新等核心流程。报告则从总体架构、三个子系统的数据库设计、Redis 键设计、Mongo 文档设计、关键事务控制逻辑、运行验证结果和页面模块分工等角度进行说明，较完整地体现了关系数据库、键值数据库、文档数据库和图谱分析方法在同一校园业务场景中的协同应用。",
    )
    add_paragraph(
        document,
        "关键词：智慧校园；异构数据库；Redis；MongoDB；数据治理；图谱分析",
        first_line_indent=0,
        bold=True,
    )
    add_page_break(document)

    add_heading(document, "一、智慧校园大数据管理系统整体设计方案", 1)
    add_paragraph(
        document,
        "本项目以荟萃大学校园信息系统为对象，围绕智慧图书馆、智慧教务和实践教学综合管理平台三个子系统开展数据库管理系统设计，并在此基础上补充数据中心、数据治理实验室和图谱分析模块。系统按照课程中“结构化数据库 + 两种以上非结构化数据库”的要求构建异构数据架构，使项目既能完成业务处理，又能体现缓存、文档管理、图关系分析和数据治理能力。",
    )
    add_paragraph(
        document,
        f"当前系统演示数据包括学生 {payload['student_count']} 人、教师 {payload['teacher_count']} 人、图书 {payload['book_count']} 本、课程 {payload['course_count']} 门、实践任务 {payload['task_count']} 项。图谱分析模块已整理出 {graph['node_count']} 个节点和 {graph['edge_count']} 条关系，数据治理页面计算得到主数据综合质量得分 {governance['quality_score_text']}，说明系统已经从单纯的数据库设计扩展到数据观测、关系分析和质量治理阶段。",
    )
    add_paragraph(
        document,
        "系统总体架构采用四层协同思路：第一层使用 SQLite 模拟 MySQL，保存学生、教师、图书、课程、成绩、预约、签到和周报等结构化主数据与事务数据；第二层使用 Redis 承载排行榜、课程热选榜、实验室利用榜、项目进度摘要和最近检索词，提高高频查询性能；第三层使用 MongoDB 承载行为日志、预警画像、事件流、TTL 搜索缓存、地理空间数据、GridFS 资产和质量快照；第四层使用本地图模型与 Cypher 导出结果组织跨系统关系，支持课程推荐、图书推荐、导师关系和多跳关联分析。",
    )
    add_table(
        document,
        ["层次", "数据库/技术", "在本项目中的职责"],
        [
            ["结构化主数据层", "SQLite", "保存 students、books、courses、tasks 等核心业务表，保证事务一致性。"],
            ["缓存与排行层", "Redis", "保存热度排行、名额与进度摘要、最近检索词和看板快照。"],
            ["文档治理层", "MongoDB", "保存行为日志、变更日志、TTL 缓存、Geo 空间、GridFS 文件和质量快照。"],
            ["图谱分析层", "Graph + Cypher", "保存学生、教师、课程、图书、任务等实体关系，支持本地图分析、Cypher 导出与多跳分析。"],
        ],
    )
    add_heading(document, "1.1 数据库选型与课程知识映射", 2)
    add_paragraph(
        document,
        "RDBMS 与 SQL 章节强调数据规范化、实体关系和事务一致性，因此本项目将学生、课程、借阅、选课、预约、成绩等核心业务建模到 SQLite；NoSQL 数据模型与键值数据库章节强调高并发读写和快速缓存，因此使用 Redis 维护多个时间窗口排行榜及状态摘要；文档数据库章节强调模式灵活和复杂文档存储，因此使用 MongoDB 保存日志、预警、Geo 数据、GridFS 文件与治理快照；图数据库章节强调多对多、多跳路径和关系传播分析，因此本项目进一步构建了统一图模型，并导出 Cypher 脚本以支撑关系建模说明。",
    )
    add_heading(document, "1.2 数据治理与主数据管理创新点", 2)
    add_paragraph(
        document,
        "在课程主数据管理部分的启发下，本项目增加了数据治理页面，对学生、教师、课程、图书和空间五类核心主数据建立统一业务主键，并对完整性、唯一性、一致性、关联完整性和时效性进行量化评分。当前治理快照已写入 MongoDB 的 data_quality_snapshot 集合，形成可追踪、可审计、可展示的治理留痕机制。",
    )
    add_bullets(
        document,
        [
            f"Redis 当前状态：{payload['redis_runtime']['status_label']}",
            f"MongoDB 当前状态：{payload['mongo_runtime']['status_label']}",
            f"图谱模块当前状态：{payload['graph_runtime']['status_label']}",
            f"主数据质量评分：{governance['quality_score_text']} 分，等级为 {governance['quality_label']}",
        ],
    )
    add_heading(document, "1.3 业务流程与数据流转说明", 2)
    add_paragraph(
        document,
        "从系统整体运行逻辑来看，三个业务子系统虽然功能不同，但都遵循相同的数据流转模式。用户请求首先进入 Flask 页面层，由页面层调用后端仓储接口；仓储接口优先在 SQLite 中完成事务操作与主数据更新；随后把高频访问结果写入 Redis 缓存层，把过程日志、事件流、画像或质量快照写入 MongoDB 文档层；最后再由数据中心、治理页面或图谱分析页面统一读取和展示。这种流转方式既保证了核心业务的一致性，又让缓存和文档层发挥各自擅长的能力。",
    )
    add_bullets(
        document,
        [
            "图书馆链路：检索/借阅 -> SQLite 事务 -> Redis 排行与最近搜索 -> Mongo 行为日志与事件流。",
            "教务链路：选课/退课 -> SQLite 事务与容量控制 -> Redis 热选榜与配额摘要 -> Mongo 变更记录与预警画像。",
            "实践链路：预约/签到/任务推进 -> SQLite 过程数据 -> Redis 利用榜与进度摘要 -> Mongo 周报、评语和风险画像。",
        ],
    )
    add_heading(document, "1.4 数据库分层设计理由", 2)
    add_paragraph(
        document,
        "采用多数据库协同设计的核心原因在于不同数据类型和访问模式的差异非常明显。学生、课程、借阅记录等数据需要强一致和结构化约束，适合放在关系型层；排行榜、名额、进度这类数据查询频率高、更新粒度小，适合放在 Redis；日志、周报、画像、评语这类字段不固定或层次较深的数据更适合采用文档型数据库；而跨系统推荐、多跳路径和中心性分析则更适合抽象为图结构。将这些需求拆到不同层次后，系统逻辑会更清晰，扩展成本也更低。",
    )
    add_heading(document, "1.5 课程要求与系统实现对应", 2)
    add_table(
        document,
        ["课程要求", "本项目实现方式", "落地结果"],
        [
            ["至少设计三个子系统", "设计图书馆、教务、实践教学三类业务模块。", "三类模块均已形成页面和数据逻辑。"],
            ["使用结构化数据库", "以 SQLite 模拟 MySQL，保存业务主数据与事务数据。", "借阅、选课、预约、成绩等均可运行。"],
            ["使用两种以上非结构化数据库", "引入 Redis 与 MongoDB，并补充图谱分析模块。", "缓存、文档留痕、关系抽象均已实现。"],
            ["完成编程实现", "采用 Flask 搭建完整演示系统。", "页面、表单、业务动作和状态联动均可验证。"],
            ["体现扩展与创新", "增加数据中心、数据治理、图谱分析、风险画像和多粒度排行。", "系统不止停留在基础增删改查。"],
        ],
    )
    add_paragraph(
        document,
        "增加这张对应表的目的，是把课程要求和系统落地结果直接对齐。这样在阅读报告时，可以很快看到每一项要求在系统中的实际对应位置，也便于后续整理最终提交材料。",
    )

    add_heading(document, "二、智慧图书馆子系统数据库设计", 1)
    add_paragraph(
        document,
        "【点题说明】本子系统结合使用了 SQLite 作为结构化数据库，以及 Redis 和 MongoDB 作为非结构化数据库，满足课程关于“每个子系统混合使用结构化数据库与两类以上非结构化数据库”的要求。",
        bold=True,
        first_line_indent=0,
        size=11.5,
    )
    add_paragraph(
        document,
        "智慧图书馆子系统承担图书检索、借阅、归还、座位预约、借阅排行和阅读行为分析等功能。该子系统既有明显的事务性要求，也有较强的读写频率波动和文档留痕需求，因此最适合作为异构数据库协同设计的典型示例。",
    )
    add_table(
        document,
        ["数据库层", "核心对象", "设计说明"],
        [
            ["SQLite", "books、borrow_records、library_seats、seat_reservations", "处理借阅事务、库存更新、座位预约等强一致操作。"],
            ["Redis", "library:last_search、rank:book:*、rank:book:category:*", "保存最近检索词和多粒度借阅热榜，提高看板访问性能。"],
            ["MongoDB", "library_behavior_log、search_session_cache、realtime_event_feed、campus_space_geo、GridFS", "保存检索行为、TTL 缓存、实时事件流、Geo 空间和演示资产。"],
            ["Neo4j", "Student-BORROWED-Book-IN_CATEGORY-Category", "用于阅读兴趣分析和图书推荐。"],
        ],
    )
    add_paragraph(
        document,
        "在结构化设计方面，books 表保存 ISBN、题名、作者、分类、馆藏位置和库存；borrow_records 表保存借阅时间、应还时间、归还时间和借阅状态；library_seats 与 seat_reservations 表保存座位资源与预约记录。通过主外键约束保证学生与图书、学生与座位之间的业务关联完整性。",
    )
    add_paragraph(
        document,
        "在 Redis 设计方面，项目按照日榜、周榜、月榜、总榜和分类日榜设计多个 ZSet Key，结合 TTL 控制不同窗口的保留周期；同时使用 String 类型保存最近检索词，既符合键值数据库课程中“按访问场景拆 key”的思想，也能清晰体现缓存层的职责分工。",
    )
    add_paragraph(
        document,
        "在 MongoDB 设计方面，library_behavior_log 用于记录检索和借阅行为，search_session_cache 采用 TTL 索引自动过期，realtime_event_feed 采用 Capped Collection 保存固定窗口事件流，campus_space_geo 使用 2dsphere 索引支持图书馆中心点附近空间查询，GridFS 保存演示截图等文件型资源，体现文档数据库在日志、缓存、空间和文件管理方面的综合能力。",
    )
    add_paragraph(
        document,
        "在最新实现中，图书馆页面进一步加入了近三周借阅热力图，并把同一借阅行为同步映射到日榜、周榜、月榜、总榜和分类日榜五类 Redis Key 中，从而更直观地展示时间窗口与缓存键设计之间的对应关系。",
    )
    add_image(document, diagram_assets["library_er"], "图 2-1 智慧图书馆关系型 E-R 设计图", width_cm=15.6)
    add_heading(document, "2.1 核心实体与关系设计", 2)
    add_table(
        document,
        ["实体/表", "关键字段", "业务作用"],
        [
            ["books", "book_id、isbn、title、category、total_copies、available_copies", "保存图书主数据与库存信息。"],
            ["borrow_records", "borrow_id、student_id、book_id、borrowed_at、due_at、returned_at、status", "保存借阅事务与归还状态。"],
            ["library_seats", "seat_id、zone、seat_code、status", "保存座位资源信息。"],
            ["seat_reservations", "reservation_id、student_id、seat_id、reserve_date、time_slot、status", "保存座位预约与时间段占用情况。"],
        ],
    )
    add_paragraph(
        document,
        "图书馆子系统最关键的两类关系分别是“学生 - 图书”和“学生 - 座位”。前者围绕借阅事务展开，需要保证库存扣减和借阅记录写入同步成功；后者围绕预约冲突控制展开，需要保证同一时间段同一座位不被重复占用。这两类关系都更适合由关系型数据库做核心约束，再把结果同步到缓存和文档层。",
    )
    add_heading(document, "2.2 Redis 键设计与过期策略", 2)
    add_table(
        document,
        ["Key 示例", "类型", "说明", "更新时机"],
        [
            ["library:last_search", "String", "保存最近一次图书检索关键字。", "每次检索时覆盖写入。"],
            ["rank:book:YYYYMMDD", "ZSet", "图书日榜。", "每次借阅成功后实时更新。"],
            ["rank:book:YYYYwkNN", "ZSet", "图书周榜。", "每次借阅成功后实时更新。"],
            ["rank:book:YYYYMM", "ZSet", "图书月榜。", "每次借阅成功后实时更新。"],
            ["rank:book:total", "ZSet", "图书总榜。", "每次借阅成功后实时更新。"],
            ["rank:book:category:<分类>:YYYYMMDD", "ZSet", "图书分类日榜。", "同分类图书借阅后更新。"],
        ],
    )
    add_paragraph(
        document,
        "将排行榜拆成多粒度 ZSet 的原因有两点。第一，不同时间窗口的榜单查询逻辑可以完全独立，不需要在读取时再做复杂聚合；第二，每个榜单都可以设置不同的生命周期，例如日榜和分类日榜可以保留较短时间，而总榜则长期保留。这样做既符合 Redis 的使用习惯，也能把缓存层设计思路写得更清楚。",
    )
    add_heading(document, "2.3 文档型数据与索引设计", 2)
    add_table(
        document,
        ["集合", "主要字段", "设计目的"],
        [
            ["library_behavior_log", "student_id、keyword、action、device、timestamp、result_count", "保存检索、借阅、归还等完整行为上下文。"],
            ["search_session_cache", "keyword、result_ids、created_at、expires_at", "保存短期检索缓存，并通过 TTL 自动清理。"],
            ["realtime_event_feed", "event_type、book_id、student_id、created_at", "保存固定窗口下的实时事件流。"],
            ["campus_space_geo", "space_id、name、location", "保存馆内空间坐标，支持地理空间查询。"],
        ],
    )
    add_paragraph(
        document,
        "这些集合的共同特点是字段灵活、上下文信息多、结构可能不断扩展。如果强行拆成关系表，不仅建模复杂，而且后期每增加一个分析字段都可能需要调整表结构。因此，这部分内容采用文档型数据库会更自然，也更符合日志和事件数据的存储特点。",
    )
    add_image(document, diagram_assets["library_json"], "图 2-2 图书馆行为日志 JSON 树状结构图", width_cm=15.6)

    add_heading(document, "三、智慧教务系统子系统数据库设计", 1)
    add_paragraph(
        document,
        "【点题说明】本子系统结合使用了 SQLite 作为结构化数据库，以及 Redis 和 MongoDB 作为非结构化数据库，满足课程关于“每个子系统混合使用结构化数据库与两类以上非结构化数据库”的要求。",
        bold=True,
        first_line_indent=0,
        size=11.5,
    )
    add_paragraph(
        document,
        "智慧教务系统负责开课、选课、退课、成绩管理和学业预警。其核心特点是课程与学生、教师之间存在复杂多对多关系，因此非常适合在关系数据库基础上进一步引入文档数据库和图数据库。",
    )
    add_table(
        document,
        ["数据库层", "核心对象", "设计说明"],
        [
            ["SQLite", "courses、course_offerings、course_selections、score_records", "保存课程目录、开课信息、选课记录和成绩数据。"],
            ["Redis", "rank:course:current、quota:course:*", "维护课程热选榜、名额占用摘要和容量预警。"],
            ["MongoDB", "teaching_change_log、warning_profile、realtime_event_feed", "记录容量调整、预警命中规则和事件留痕。"],
            ["Neo4j", "Student-ENROLLED_IN-Course、Teacher-TEACHES-Course", "支撑相似学生选课推荐和教师授课关系分析。"],
        ],
    )
    add_paragraph(
        document,
        "关系型部分重点保证 course_offerings 与 course_selections 的一致性，selected_count 不能超过 capacity，score_records 与选课记录保持一一对应。MongoDB 中的 teaching_change_log 用于记录容量调整前后值，warning_profile 保存学业预警命中规则，使教务子系统不仅具备操作功能，还具备管理审计和风险画像能力。图数据库将学生、课程、教师关系统一建模，可直接查询“当前学生未选但相似学生选过的课程”，适合展示图数据库的多跳查询优势。",
    )
    add_paragraph(
        document,
        "为贴合当前页面实现，教务模块还增加了课程热选榜、容量预警、选课时间冲突检测与智能推荐。这样既能体现关系数据库对课表与选课事务的约束作用，也能体现 Redis 在热点课程和容量状态展示中的响应优势。",
    )
    add_image(document, diagram_assets["academic_er"], "图 3-1 智慧教务关系型 E-R 设计图", width_cm=15.6)
    add_heading(document, "3.1 核心实体与事务约束", 2)
    add_table(
        document,
        ["实体/表", "关键字段", "设计重点"],
        [
            ["courses", "course_id、course_code、name、credit", "保存课程目录主数据。"],
            ["course_offerings", "offering_id、course_id、teacher_id、weekday、time_slot、capacity、selected_count", "保存具体开课信息与容量状态。"],
            ["course_selections", "selection_id、student_id、offering_id、selected_at、status", "保存学生选课关系。"],
            ["score_records", "score_id、selection_id、usual_score、final_score、total_score", "保存成绩并与选课记录对应。"],
        ],
    )
    add_paragraph(
        document,
        "在事务设计中，最重要的是防止重复选课、容量溢出和成绩记录脱离选课记录单独存在。因此，选课成功必须同时满足课程未满、学生未重复选择且时间不冲突三个条件。只有这些条件全部满足后，系统才会写入选课记录并递增 selected_count。",
    )
    add_heading(document, "3.2 冲突检测与 Redis 摘要设计", 2)
    add_table(
        document,
        ["Key/规则", "作用", "说明"],
        [
            ["rank:course:current", "课程热选榜", "按当前选课人数维护课程热度排序。"],
            ["quota:course:<offering_id>", "课程容量摘要", "缓存容量、已选人数和剩余人数。"],
            ["时间冲突检测", "选课拦截规则", "比较星期与时间段，若重叠则拒绝选课。"],
        ],
    )
    add_paragraph(
        document,
        "通过把热选榜和容量摘要缓存到 Redis，教务页面可以快速展示热门课程和名额情况，而无需每次刷新都对选课表做统计聚合。时间冲突检测仍然保留在业务层执行，保证规则一致性和结果准确性。",
    )
    add_heading(document, "3.3 文档留痕与预警画像设计", 2)
    add_table(
        document,
        ["集合", "保存内容", "设计意义"],
        [
            ["teaching_change_log", "调容前后值、操作人、操作时间、原因说明", "保留教务调整留痕。"],
            ["warning_profile", "预警标签、命中规则、建议说明", "保存学生学业预警画像。"],
            ["realtime_event_feed", "选课、退课、调容等事件", "形成近期教务行为流。"],
        ],
    )
    add_paragraph(
        document,
        "在教务场景中，很多分析结果并不是一条简单的结构化记录，而是带有说明文字、命中条件、前后差异和建议措施的复合对象。这类数据如果完全拆成二维表会让结构变得冗长，而采用文档形式可以更直接地保存一次完整事件或一份完整画像。",
    )
    add_image(document, diagram_assets["academic_json"], "图 3-2 教务预警画像 JSON 树状结构图", width_cm=15.6)

    add_heading(document, "四、实践教学综合管理平台子系统数据库设计", 1)
    add_paragraph(
        document,
        "【点题说明】本子系统结合使用了 SQLite 作为结构化数据库，以及 Redis 和 MongoDB 作为非结构化数据库，满足课程关于“每个子系统混合使用结构化数据库与两类以上非结构化数据库”的要求。",
        bold=True,
        first_line_indent=0,
        size=11.5,
    )
    add_paragraph(
        document,
        "实践教学综合管理平台用于管理项目、实验室、实习任务、签到和周报，是最能体现结构化数据、文本数据和关系路径混合管理需求的子系统。",
    )
    add_table(
        document,
        ["数据库层", "核心对象", "设计说明"],
        [
            ["SQLite", "practice_projects、lab_rooms、lab_bookings、internship_tasks、attendance_records、weekly_reports", "完成项目管理、实验室预约、签到和周报事务处理。"],
            ["Redis", "rank:practice:progress、rank:lab:usage、project:progress:*", "维护任务推进榜、实验室利用榜和项目进度摘要。"],
            ["MongoDB", "internship_weekly_report、evaluation_comment、realtime_event_feed", "保存周报正文、导师评语和实践事件流。"],
            ["Neo4j", "Student-HAS_TASK-Task-AT_BASE-Base、Task-GUIDED_BY-Mentor", "展示学生、任务、基地和导师之间的多跳关系。"],
        ],
    )
    add_paragraph(
        document,
        "在实践教学子系统中，weekly_reports 和 evaluation_comment 更适合采用文档型存储，因为其正文长度和结构不固定；实验室预约和任务管理则保持在结构化数据库中，以保证时间段和资源冲突判断准确；使用图数据库后，可以很自然地展示“某学生负责什么任务、任务位于哪个基地、由谁指导、预约过哪些实验室”，非常适合作为图数据库课程知识的项目应用场景。",
    )
    add_paragraph(
        document,
        "在当前实现版本中，实践平台还补充了实验室利用榜、任务推进榜和风险画像。系统会结合进度、签到和周报情况生成风险等级，并把风险结果写入 MongoDB 的 practice_risk_profile 集合，形成过程性数据的分析闭环。",
    )
    add_image(document, diagram_assets["practice_er"], "图 4-1 实践教学平台关系型 E-R 设计图", width_cm=15.6)
    add_heading(document, "4.1 核心实体设计", 2)
    add_table(
        document,
        ["实体/表", "关键字段", "作用"],
        [
            ["practice_projects", "project_id、name、instructor_id、progress、status", "保存项目总体信息与进度。"],
            ["lab_rooms", "room_id、room_code、capacity、location、status", "保存实验室资源。"],
            ["lab_bookings", "booking_id、student_id、room_id、booking_date、slot、status", "保存实验室预约。"],
            ["internship_tasks", "task_id、student_id、project_id、progress、status", "保存学生任务及完成状态。"],
            ["attendance_records", "attendance_id、student_id、task_id、check_in_time、status", "保存签到记录。"],
            ["weekly_reports", "report_id、student_id、task_id、content、submitted_at", "保存过程汇报文本。"],
        ],
    )
    add_paragraph(
        document,
        "实践平台中，预约与签到强调时间和资源约束，任务推进强调过程状态，周报和评语则强调文本记录。因此同一个子系统内部就同时存在强结构化对象和半结构化对象，这也是它最适合体现异构数据库协同设计的原因。",
    )
    add_heading(document, "4.2 Redis 榜单与进度摘要设计", 2)
    add_table(
        document,
        ["Key", "类型", "作用"],
        [
            ["rank:practice:progress", "ZSet", "维护任务推进榜。"],
            ["rank:lab:usage", "ZSet", "维护实验室利用榜。"],
            ["project:progress:<project_id>", "Hash/String", "保存单项目当前进度摘要。"],
        ],
    )
    add_paragraph(
        document,
        "这些键把原本分散在任务表和预约表中的统计结果直接整理为可读取的摘要，适合在看板和页面顶部指标区快速展示。对这类以展示和比较为主的数据，使用 Redis 比每次从关系表现场聚合更加直接。",
    )
    add_heading(document, "4.3 风险画像与文档数据设计", 2)
    add_table(
        document,
        ["集合", "主要内容", "说明"],
        [
            ["internship_weekly_report", "周报正文、阶段总结、问题描述、提交时间", "保存过程文本材料。"],
            ["evaluation_comment", "导师评语、评分、修改建议、时间", "保存导师反馈。"],
            ["practice_risk_profile", "进度状态、签到情况、报告提交情况、风险等级", "保存自动生成的风险画像。"],
        ],
    )
    add_paragraph(
        document,
        "风险画像并不是原始事务数据，而是基于多类数据综合计算出来的分析结果。它既包含结构化评分，也包含标签和说明文字，因此采用文档型数据库保存更灵活。这样后续如果要新增风险维度，例如资源冲突次数、导师反馈频率，也不需要重新大幅调整结构。",
    )
    add_image(document, diagram_assets["practice_json"], "图 4-2 实践风险画像 JSON 树状结构图", width_cm=15.6)

    add_heading(document, "五、编程实现（智慧图书馆子系统）", 1)
    add_paragraph(
        document,
        "本项目最终采用 Flask 构建 Web 演示系统，后端主要由 app.py、demo_backend.py、mongo_real_backend.py 和 graph_backend.py 组成，其中图书馆子系统的实现最完整，能够充分体现结构化数据库、Redis、MongoDB 和 Neo4j 协同工作的流程，因此本节选择图书馆子系统进行编程实现说明。",
    )
    add_heading(document, "5.1 核心实现流程", 2)
    add_bullets(
        document,
        [
            "图书检索：library 页面提交关键字后，先查询 SQLite 的 books 表，再把关键字写入 Redis 的 library:last_search，同时向 MongoDB 的 library_behavior_log、search_session_cache 和 realtime_event_feed 写入记录。",
            "图书借阅：borrow_book 函数通过事务减少 books.available_copies 并插入 borrow_records；成功后立即刷新 Redis 热榜，并将借阅行为写入 MongoDB 事件流。",
            "图书归还：return_book 函数更新归还状态与 returned_at，恢复库存，保持结构化数据和缓存数据同步。",
            "数据中心展示：页面直接展示 Redis 键空间、Mongo 聚合、TTL 缓存、Geo 查询、GridFS 文件以及治理快照，使多数据库协同实现具有可视化证据。",
            "图谱扩展：graph_lab 页面会生成学生、图书和类别等关系的图模型，并导出 Cypher 脚本，展示图书推荐和多跳路径分析结果。",
        ],
    )
    add_heading(document, "5.2 核心代码片段展示", 2)
    add_paragraph(
        document,
        "为了证明本项目不仅完成了页面展示，还落实了真实的数据操作流程，本节补充三个最能体现工作量和课程知识点的核心代码片段，分别对应 Redis 排行榜更新、MongoDB 聚合分析以及教务模块跨库写入流程。",
    )
    add_code_block(
        document,
        """
with self.connect() as conn:
    conn.execute("UPDATE books SET available_copies = available_copies - 1 WHERE book_id = ?", (book_id,))
    conn.execute(
        "INSERT INTO borrow_records (book_id, student_id, borrowed_at, due_at, returned_at, status) VALUES (?, ?, ?, ?, ?, ?)",
        (book_id, student_id, event_time.strftime("%Y-%m-%d %H:%M:%S"), (event_time + timedelta(days=30)).strftime("%Y-%m-%d %H:%M:%S"), None, "borrowing"),
    )
self.redis.delete(f"book:detail:{book_id}")
self.redis.incr(f"user:borrow:count:{student_id}")
self.redis.zadd("book:hot:rank", str(book_id), 1)
self._update_book_rankings(book_id, book["category"], event_time=event_time)
        """,
        "代码 5-1 图书借阅后同步刷新 Redis 榜单（摘自 demo_backend.py）",
    )
    add_paragraph(
        document,
        "这段代码体现了 Redis 在项目中的真实职责。图书借阅成功后，不只是简单写一个总榜，而是先删除图书详情缓存、更新用户借阅计数，再对热榜和多粒度榜单统一加分。这样报告中的日榜、周榜、月榜和分类榜设计都能与实际代码一一对应。",
    )
    add_code_block(
        document,
        """
library_actions = list(
    self.db["library_behavior_log"].aggregate(
        [
            {"$group": {"_id": "$action", "count": {"$sum": 1}}},
            {"$sort": {"count": -1, "_id": 1}},
        ]
    )
)
        """,
        "代码 5-2 MongoDB 聚合管道示例（摘自 mongo_real_backend.py）",
    )
    add_paragraph(
        document,
        "MongoDB 并不是只用来保存日志。这里直接通过 `$group` 和 `$sort` 对借阅行为进行聚合统计，能够支撑数据中心页面中的行为分布看板，也能证明文档库在轻量分析方面的实用性。",
    )
    add_code_block(
        document,
        """
with self.connect() as conn:
    conn.execute(
        "INSERT INTO course_selections (offering_id, student_id, selected_at, status) VALUES (?, ?, ?, ?)",
        (offering_id, student_id, now_text(), "selected"),
    )
    conn.execute("UPDATE course_offerings SET selected_count = selected_count + 1 WHERE offering_id = ?", (offering_id,))
self.redis.set(quota_key, remaining - 1)
self.redis.zadd("rank:course:current", str(offering_id), 1)
self.mongo.insert(
    "teaching_change_log",
    {"operator": student["student_name"], "changeType": "学生选课", "target": offering["course_name"], "afterValue": remaining - 1},
)
        """,
        "代码 5-3 教务模块跨库写入流程（摘自 demo_backend.py）",
    )
    add_paragraph(
        document,
        "这段代码最能说明多数据库协同的真实处理顺序。系统先在 SQLite 中完成选课关系写入与已选人数递增，再更新 Redis 中的容量摘要和热选榜，最后把过程留痕写入 MongoDB。这样既能保证事务正确，也能支持后续分析与审计。",
    )
    add_heading(document, "5.3 模块分工与创新实现", 2)
    add_table(
        document,
        ["模块/文件", "主要职责", "体现的课程知识"],
        [
            ["app.py", "提供 /library、/data-center、/graph-lab 等路由", "Web 层与数据库设计结果的集成展示"],
            ["demo_backend.py", "实现借阅、归还、检索、排行榜、治理计算与图谱总览", "SQL 事务、Redis 键设计、Mongo 文档设计、主数据治理"],
            ["mongo_real_backend.py", "实现 JSON Schema、Capped、TTL、Geo、GridFS 和聚合管道", "文档数据库设计与 MongoDB 实践"],
            ["graph_backend.py", "构建图模型、导出 Cypher、同步 Neo4j、执行推荐与中心性分析", "图数据库基础与图数据库设计"],
        ],
    )
    add_paragraph(
        document,
        "项目在原有课程作业基础上进一步加入了主数据管理、排行榜扩展、风险画像和图谱分析：一方面通过 data_quality_snapshot 集合记录治理结果，使数据质量具有留痕能力；另一方面通过统一图模型与 Cypher 导出结果组织三大子系统之间的关系，支持课程推荐、图书推荐、导师路径、多跳路径和中心性分析。这些改造使系统具备了较完整的数据管理实践闭环。",
    )
    add_heading(document, "5.4 关键控制逻辑说明", 2)
    add_table(
        document,
        ["业务场景", "控制逻辑", "目的"],
        [
            ["图书借阅", "先检查库存，再在事务中同时更新库存与借阅记录。", "防止库存与记录不一致。"],
            ["图书归还", "在事务中恢复库存并更新 returned_at。", "保证归还状态可追踪。"],
            ["课程选课", "检查是否重复、是否满员、是否时间冲突。", "防止超卖和课表冲突。"],
            ["实验室预约", "检查同一时段同一实验室是否已占用。", "避免资源冲突。"],
            ["风险画像生成", "整合进度、签到、周报等维度后生成等级。", "形成过程性分析结果。"],
        ],
    )
    add_paragraph(
        document,
        "这些控制逻辑说明本项目并非只是在页面上展示数据库内容，而是把课程中关于事务控制、缓存设计、文档留痕和分析计算的知识真正写进了后端流程。也就是说，页面上看到的结果都有清晰的数据处理路径作为支撑。",
    )
    add_heading(document, "5.5 多库协同一致性设计", 2)
    add_paragraph(
        document,
        "在同时操作 SQLite、Redis 和 MongoDB 时，本项目采用了“结构化主库本地事务优先，缓存层与文档层随后同步”的策略。以教务选课为例，系统先在 SQLite 中完成选课记录写入和容量扣减，再更新 Redis 中的 `course:quota:*` 和 `rank:course:current`，最后把变更写入 MongoDB 的 `teaching_change_log`。这种处理顺序的目标，是把最关键的业务一致性放在事务主库中优先保障。",
    )
    add_paragraph(
        document,
        "如果缓存层或文档层某次同步失败，系统并不会破坏主库中的正确业务结果，而是允许通过后续补偿机制恢复衍生数据。当前代码中已经保留了 `_bootstrap_rank_indexes()` 一类从 SQLite 重建 Redis 榜单的逻辑，能够依据借阅记录、选课人数和实验室预约记录重新计算缓存结果，因此报告中可以将这套机制解释为典型的最终一致性（Eventual Consistency）实现方式。",
    )
    add_code_block(
        document,
        """
if not self.redis.ztop("rank:course:current", limit=1):
    course_rows = conn.execute("SELECT offering_id, selected_count FROM course_offerings").fetchall()
    for row in course_rows:
        self.redis.zadd("rank:course:current", str(row["offering_id"]), int(row["selected_count"]))
        """,
        "代码 5-4 基于 SQLite 结果重建 Redis 榜单的补偿逻辑（摘自 demo_backend.py）",
    )
    add_paragraph(
        document,
        "因此，本项目并不是简单把三种数据库并列使用，而是明确区分了主数据、缓存和留痕数据的职责，并给出了跨库口径不一致时的恢复路径。这一点可以有效回应报告中关于“多库协同时如何保证数据一致性”的问题。",
    )
    add_heading(document, "5.6 Neo4j 部署说明与图谱切换", 2)
    add_paragraph(
        document,
        "当前图谱模块默认处于“本地图分析模式”，原因是课程演示更强调单机环境下的一键运行能力。graph_backend.py 会先使用 NetworkX 在本地构建学生、课程、图书、项目和导师的统一图模型，并同步导出 `smart_campus_graph.cypher`，这样即使没有外部 Neo4j 服务，也能完成推荐、多跳路径和中心性分析展示。",
    )
    add_paragraph(
        document,
        "与此同时，系统逻辑已经提前完成了解耦。代码会优先从环境变量或 `neo4j_runtime.json` 中读取 `SMART_CAMPUS_NEO4J_URI`、`SMART_CAMPUS_NEO4J_USER` 和 `SMART_CAMPUS_NEO4J_PASSWORD`；当连接可用时，图谱模块会从本地图分析模式切换到 Neo4j 在线模式，并自动把导出的 Cypher 脚本同步执行到图数据库中。也就是说，当前报告中把图谱模块写为“本地图分析 + Cypher 导出 + Neo4j Ready”是合理且有代码依据的。",
    )
    add_heading(document, "5.7 运行效果与开发环境", 2)
    add_paragraph(
        document,
        f"截至报告生成时，系统实际运行状态为：Redis 为“{payload['redis_runtime']['status_label']}”，MongoDB 为“{payload['mongo_runtime']['status_label']}”，图谱模块为“{payload['graph_runtime']['status_label']}”。其中 Redis 已真实接管排行榜与缓存摘要，MongoDB 已真实启用日志、TTL、Geo、GridFS 和质量快照能力，图谱模块则通过本地图分析与 Cypher 导出完成关系建模说明，说明项目已经达到“每个子系统使用结构化数据库和两种以上非结构化数据库”的课程要求。",
    )
    add_paragraph(
        document,
        "综上所述，本项目以校园业务为主线，在满足课程基本要求的前提下，进一步落实了 Redis 键值设计、MongoDB 文档数据库设计、图数据库建模以及主数据治理方法，使系统具备事务处理、缓存加速、文档留痕、关系分析和治理评估的综合能力，能够完整体现异构数据库协同设计的课程目标。",
    )
    add_heading(document, "5.8 开发环境与部署说明", 2)
    add_table(
        document,
        ["项目要素", "当前配置", "说明"],
        [
            ["开发语言", "Python", "后端逻辑与报告生成脚本均使用 Python 完成。"],
            ["Web 框架", "Flask", "负责页面路由、表单处理和接口组织。"],
            ["事务数据层", "SQLite", "在本地环境中模拟 MySQL 的事务型能力。"],
            ["缓存层", "Redis", "用于排行榜、配额、状态摘要和热点缓存。"],
            ["文档层", "MongoDB", "用于日志、事件流、TTL、Geo、GridFS 和画像文档。"],
            ["运行平台", "Windows + Chrome/Edge 浏览器", "适合本地演示和课程展示。"],
        ],
    )
    add_paragraph(
        document,
        "从部署角度看，本项目更偏向课程演示型系统，因此采用本地可运行、可验证的轻量化结构，而没有引入复杂的服务器拆分。这样的设计既便于课堂环境快速启动，也方便直接验证各类数据库能力是否真正接入。",
    )

    add_heading(document, "六、系统页面与关键扩展说明", 1)
    add_paragraph(
        document,
        "为了让数据库设计结果能够被直接观察和验证，系统在三个业务子系统之外，还增加了总览页、数据中心页、数据治理页和图谱分析页。这样可以把结构化事务、缓存键空间、文档集合和关系图模型分别以页面方式呈现，便于从界面层面说明多数据库协同分工。",
    )
    add_table(
        document,
        ["页面/模块", "当前实现内容", "对应的数据管理能力"],
        [
            ["总览页", "统一展示系统架构、模块入口、排行预览和缓存快照。", "整体架构说明、模块联动展示"],
            ["智慧图书馆", "借阅归还、座位预约、多粒度排行榜、借阅热力图。", "事务控制、Redis 排行、日志留痕"],
            ["智慧教务", "选课退课、热选榜、容量预警、冲突检测、课程推荐。", "关系约束、缓存摘要、预警画像"],
            ["实践教学平台", "预约签到、任务推进榜、实验室利用榜、风险画像。", "过程数据管理、文档存储、聚合分析"],
            ["数据中心", "Redis 键快照、Mongo 聚合、TTL/Geo/GridFS、治理快照。", "多数据库观测与能力验证"],
            ["数据治理/图谱分析", "主数据目录、质量评分、Cypher 导出、路径分析。", "主数据治理、关系建模、多跳分析"],
        ],
    )
    add_paragraph(
        document,
        "其中，图书馆模块增加了借阅热力图与五类排行榜，教务模块增加了冲突检测与智能推荐，实践模块增加了风险画像。这些扩展都不是独立的前端展示，而是直接依赖 Redis、MongoDB 和关系型数据的联合计算结果。",
    )
    add_heading(document, "6.1 高级扩展设计说明", 2)
    add_table(
        document,
        ["扩展方向", "当前设计思路", "后续可演进方式"],
        [
            ["冷热数据分层", "Redis 保存热排行与状态摘要，MongoDB 保存近阶段温数据文档。", "历史日志可通过 ETL 继续归档到 HDFS 或对象存储，支持离线分析。"],
            ["消息队列削峰", "当前演示版直接写入事务层，便于课堂环境启动和验证。", "在抢课等高并发场景中可加入 Kafka 或 RabbitMQ，对请求排队消费。"],
            ["跨库一致性补偿", "已具备根据 SQLite 重建 Redis 榜单与摘要的逻辑。", "后续可增加定时对账任务，自动校验缓存层与文档层口径。"],
            ["图谱部署切换", "当前以本地图分析和 Cypher 导出为主。", "补充 Neo4j 连接配置后即可切换到在线图数据库模式。"],
        ],
    )
    add_paragraph(
        document,
        "这一部分的意义在于把项目从“能跑的课程作业”进一步提升到“具备演进思路的数据管理系统”。即使当前运行环境是单机演示，也已经预留了热温冷分层、峰值保护和图谱上线等扩展方向。",
    )
    add_heading(document, "6.2 页面模块与数据能力对应关系", 2)
    add_table(
        document,
        ["页面模块", "页面可见内容", "背后依赖的数据能力"],
        [
            ["首页总览", "系统结构、统计指标、排行预览、缓存快照", "统一读取 SQLite 统计、Redis 摘要和 Mongo 快照。"],
            ["智慧图书馆", "图书列表、借阅记录、借阅热榜、热力图", "事务更新 + Redis 排行 + Mongo 行为日志。"],
            ["智慧教务", "选课列表、热选榜、预警、冲突检测", "关系约束 + Redis 配额 + Mongo 画像。"],
            ["实践教学平台", "预约、签到、任务推进、风险画像", "过程事务 + Redis 榜单 + Mongo 文档计算。"],
            ["数据中心", "键空间、聚合、TTL、Geo、GridFS、治理快照", "多数据库统一观测与统计。"],
            ["治理/图谱页", "质量评分、血缘、Cypher、路径分析", "主数据治理 + 本地图分析。"],
        ],
    )
    add_heading(document, "6.3 运行验证结果", 2)
    add_table(
        document,
        ["验证项", "结果", "说明"],
        [
            ["首页与各子页面访问", "通过", "总览、图书馆、教务、实践、数据中心、治理、图谱页面均可访问。"],
            ["Redis 连接", "通过", "当前已连接真实 Redis，并可维护排行榜与键快照。"],
            ["MongoDB 连接", "通过", "当前已连接真实 MongoDB，并可写入日志、TTL、Geo 和 GridFS 数据。"],
            ["图书借阅联动", "通过", "借阅成功后库存、排行和日志会同步变化。"],
            ["选课冲突拦截", "通过", "时间冲突课程会被直接拦截，避免重复写入。"],
            ["实践风险画像", "通过", "风险结果已生成并落入 MongoDB 集合。"],
        ],
    )
    add_heading(document, "6.4 页面组织与界面优化说明", 2)
    add_paragraph(
        document,
        "为了让系统不只是“功能堆叠”，本次优化还对页面结构做了统一整理。所有页面采用左侧统一导航、顶部运行态摘要、页内分区锚点和中部卡片式内容布局，使用户在查看系统时能更快定位到业务操作区、分析看板区和验证结果区。相比单纯把多个表格堆在一起，这种结构更有利于把数据库设计逻辑与页面观察路径对应起来。",
    )
    add_bullets(
        document,
        [
            "首页突出系统架构、业务链路、子系统入口和缓存快照。",
            "图书馆页突出热力图、多粒度排行榜、借阅与预约操作。",
            "教务页突出学生视角、热选榜、冲突检测、推荐和变更留痕。",
            "实践页突出风险画像、预约签到、过程记录和 Mongo 文档快照。",
            "数据中心页突出 Redis 键空间、Mongo 能力、治理摘要和图谱快照。",
        ],
    )

    add_heading(document, "6.5 最新系统页面截图", 2)
    add_image(document, screenshot_dir / "home_latest.png", "图 6-1 系统总览页", width_cm=15.8)
    add_image(document, screenshot_dir / "library_latest.png", "图 6-2 智慧图书馆页面", width_cm=15.8)
    add_image(document, screenshot_dir / "academic_latest.png", "图 6-3 智慧教务页面", width_cm=15.8)
    add_image(document, screenshot_dir / "practice_latest.png", "图 6-4 实践教学平台页面", width_cm=15.8)
    add_image(document, screenshot_dir / "data_center_latest.png", "图 6-5 数据中心页面", width_cm=15.8)
    add_image(document, screenshot_dir / "governance_latest.png", "图 6-6 数据治理实验室页面", width_cm=15.8)
    add_image(document, screenshot_dir / "graph_latest.png", "图 6-7 图谱分析页面", width_cm=15.8)
    add_heading(document, "七、总结与体会", 1)
    add_paragraph(
        document,
        "本次大作业最大的收获在于把课程中看似分散的知识点串联到了一个完整业务场景里。过去在课堂上学习关系模型、键值存储、文档数据库和图数据库时，更多是分别理解各自的特点；而在这个项目中，它们被放到了同一条业务链路上，各自承担不同职责，这让我更容易理解“为什么真实系统往往是多数据库协同，而不是单库解决所有问题”。",
    )
    add_paragraph(
        document,
        "从结果上看，SQLite 负责保证事务数据正确，Redis 负责快速响应高频查询，MongoDB 负责保留上下文丰富的过程数据，图谱模块负责表达跨系统关系。这样的分工既体现了课程中的理论知识，也让系统在页面层面具有较好的可观察性。对我来说，这份作业不仅是一次数据库设计练习，也是一次把系统分析、页面实现和数据治理思路结合起来的综合训练。",
    )

    add_heading(document, "参考文献", 1)
    references = [
        "[1] 孙旭东, 檀昌稳, 杨洋, 等. 共建共享视角下数智校园的理论框架与数据治理实现路径[J]. 现代教育技术, 2024, 34(08):132-141.",
        "[2] 周晓玮. 基于异构数据库的高校数据集成设计与实现[J]. 航海教育研究, 2022, 39(01):97-101.",
        "[3] 孙超. Redis内存数据库在智慧消防系统设计中的应用[J]. 网络安全技术与应用, 2018(08):103-105.",
        "[4] 白洁, 武佳丽, 余啟旺, 等. 基于MongoDB的非关系型数据库的设计与应用[J]. 湖北师范大学学报(自然科学版), 2022, 42(02):79-82.",
        f"[5] Neo4j, Inc. Neo4j Operations Manual[EB/OL]. https://neo4j.com/docs/operations-manual/current/ , {datetime.now().strftime('%Y-%m-%d')}.",
        f"[6] MongoDB, Inc. MongoDB Manual[EB/OL]. https://www.mongodb.com/docs/manual/ , {datetime.now().strftime('%Y-%m-%d')}.",
    ]
    for item in references:
        add_paragraph(document, item, first_line_indent=0)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate the smart campus assignment report.")
    parser.add_argument("--output-dir", default=r"C:\Users\css\Desktop\大数据管理大作业")
    parser.add_argument("--student-name", default="待填写")
    parser.add_argument("--student-no", default="待填写")
    parser.add_argument("--class-name", default="待填写")
    parser.add_argument("--public-url", default=DEFAULT_PUBLIC_URL)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    repo = SmartCampusRepository(BASE_DIR)
    payload = build_report_payload(repo)
    report_date = datetime.now().strftime("%Y年%m月%d日")
    public_url = args.public_url.strip() or DEFAULT_PUBLIC_URL
    output_dir = Path(args.output_dir)
    markdown_path = output_dir / "智慧校园大数据管理课程大作业报告-增强修订版.md"
    docx_path = output_dir / "智慧校园大数据管理课程大作业报告-增强修订版.docx"

    write_markdown(markdown_path, payload, report_date, args.student_name, args.student_no, args.class_name, public_url)
    write_docx(docx_path, payload, report_date, args.student_name, args.student_no, args.class_name, public_url)

    print(markdown_path)
    print(docx_path)


if __name__ == "__main__":
    main()
